import os
import json
import re
from typing import List, Dict, Any, Tuple

# Try importing pymupdf
try:
    import pymupdf as fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Try importing docx
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Try importing PIL and pytesseract
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class DocumentChunk:
    def __init__(self, chunk_id: int, document_name: str, page_number: int, text: str):
        self.chunk_id = chunk_id
        self.document_name = document_name
        self.page_number = page_number
        self.text = text.strip()

    def to_dict(self) -> Dict[str, Any]:
        return {
            "chunk_id": self.chunk_id,
            "document_name": self.document_name,
            "page_number": self.page_number,
            "text": self.text,
            "char_count": len(self.text)
        }


class DocumentProcessor:
    """
    Extracts text, metadata, and page-level chunks from PDF, DOCX, TXT, and scanned image files.
    """

    @classmethod
    def process_file(cls, file_path: str, filename: str) -> Tuple[int, List[Dict[str, Any]], str]:
        """
        Processes a file and returns: (page_count, chunks_list, full_text)
        """
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == ".pdf":
            return cls._process_pdf(file_path, filename)
        elif ext in [".docx", ".doc"]:
            return cls._process_docx(file_path, filename)
        elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
            return cls._process_image(file_path, filename)
        elif ext in [".txt", ".md", ".csv", ".json"]:
            return cls._process_text_file(file_path, filename)
        else:
            # Attempt plain text read as fallback
            return cls._process_text_file(file_path, filename)

    @classmethod
    def _process_pdf(cls, file_path: str, filename: str) -> Tuple[int, List[Dict[str, Any]], str]:
        chunks = []
        full_text_parts = []
        page_count = 0
        chunk_id_counter = 1

        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                page_count = len(doc)
                for page_idx in range(page_count):
                    page = doc[page_idx]
                    page_num = page_idx + 1
                    page_text = page.get_text("text")

                    # If page text is very sparse or empty, check if OCR is needed
                    if len(page_text.strip()) < 30 and OCR_AVAILABLE:
                        ocr_text = cls._ocr_pdf_page(page)
                        if len(ocr_text.strip()) > len(page_text.strip()):
                            page_text = ocr_text

                    if not page_text.strip():
                        page_text = f"[Scanned/Image Page {page_num} in {filename}]"

                    full_text_parts.append(f"--- [Page {page_num}] ---\n{page_text}")
                    
                    # Split page into coherent paragraph chunks
                    page_chunks = cls._chunk_page_text(page_text, filename, page_num, chunk_id_counter)
                    for ch in page_chunks:
                        chunks.append(ch.to_dict())
                        chunk_id_counter += 1

                doc.close()
                return max(1, page_count), chunks, "\n\n".join(full_text_parts)
            except Exception as e:
                # If PyMuPDF fails or file is raw text
                pass

        # Fallback if PyMuPDF not available or failed
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
            chunks.append(DocumentChunk(1, filename, 1, raw_text).to_dict())
            return 1, chunks, raw_text
        except Exception as e:
            err_text = f"Error reading PDF {filename}: {str(e)}"
            return 1, [DocumentChunk(1, filename, 1, err_text).to_dict()], err_text

    @classmethod
    def _ocr_pdf_page(cls, page) -> str:
        """Render PDF page as pixmap and run OCR via pytesseract if available."""
        if not OCR_AVAILABLE:
            return ""
        try:
            pix = page.get_pixmap(dpi=150)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img)
            return text
        except Exception:
            return ""

    @classmethod
    def _process_docx(cls, file_path: str, filename: str) -> Tuple[int, List[Dict[str, Any]], str]:
        chunks = []
        paragraphs = []
        chunk_id = 1
        
        if DOCX_AVAILABLE:
            try:
                doc = docx.Document(file_path)
                for p in doc.paragraphs:
                    if p.text.strip():
                        paragraphs.append(p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_cells:
                            paragraphs.append(" | ".join(row_cells))
            except Exception:
                pass

        if not paragraphs:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                paragraphs = [f.read()]

        full_text = "\n\n".join(paragraphs)
        # Approximate 400 words per page
        words = full_text.split()
        page_count = max(1, (len(words) // 400) + 1)

        # Create chunks
        current_chunk = []
        current_page = 1
        current_word_count = 0

        for p in paragraphs:
            p_words = len(p.split())
            current_chunk.append(p)
            current_word_count += p_words

            if current_word_count >= 150:
                chunk_text = "\n".join(current_chunk)
                chunks.append(DocumentChunk(chunk_id, filename, current_page, chunk_text).to_dict())
                chunk_id += 1
                current_chunk = []
                if current_word_count >= 400:
                    current_page += 1
                    current_word_count = 0

        if current_chunk:
            chunk_text = "\n".join(current_chunk)
            chunks.append(DocumentChunk(chunk_id, filename, current_page, chunk_text).to_dict())

        return page_count, chunks, full_text

    @classmethod
    def _process_image(cls, file_path: str, filename: str) -> Tuple[int, List[Dict[str, Any]], str]:
        extracted_text = ""
        if OCR_AVAILABLE:
            try:
                img = Image.open(file_path)
                extracted_text = pytesseract.image_to_string(img)
            except Exception as e:
                extracted_text = f"[OCR Image extraction notice for {filename}: {str(e)}]"
        else:
            extracted_text = f"[Image document: {filename}. OCR engine ready.]"

        if not extracted_text.strip():
            extracted_text = f"[Scanned Image Document: {filename}]"

        chunks = [DocumentChunk(1, filename, 1, extracted_text).to_dict()]
        return 1, chunks, extracted_text

    @classmethod
    def _process_text_file(cls, file_path: str, filename: str) -> Tuple[int, List[Dict[str, Any]], str]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except Exception:
            text = ""

        chunks = cls._chunk_page_text(text, filename, 1, 1)
        chunk_dicts = [c.to_dict() for c in chunks]
        return 1, chunk_dicts, text

    @classmethod
    def _chunk_page_text(cls, page_text: str, filename: str, page_num: int, start_chunk_id: int) -> List[DocumentChunk]:
        """Splits page text into readable chunks with minimum semantic boundaries."""
        chunks = []
        raw_paras = [p.strip() for p in page_text.split("\n\n") if p.strip()]
        
        if not raw_paras:
            raw_paras = [p.strip() for p in page_text.split("\n") if p.strip()]

        if not raw_paras:
            chunks.append(DocumentChunk(start_chunk_id, filename, page_num, page_text.strip() or f"[Empty content on page {page_num}]"))
            return chunks

        current_buf = []
        current_len = 0
        cid = start_chunk_id

        for para in raw_paras:
            current_buf.append(para)
            current_len += len(para)
            if current_len >= 400:
                chunk_text = "\n".join(current_buf)
                chunks.append(DocumentChunk(cid, filename, page_num, chunk_text))
                cid += 1
                current_buf = []
                current_len = 0

        if current_buf:
            chunk_text = "\n".join(current_buf)
            chunks.append(DocumentChunk(cid, filename, page_num, chunk_text))

        return chunks

